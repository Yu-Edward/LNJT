# import socket
# import sqlite3
# import time
# import os
# import datetime
# import configparser
# import xlrd
# import openpyxl
# import winreg
# import zipfile
import socket,sqlite3,time,os,datetime,configparser,xlrd,openpyxl,winreg,zipfile,shutil,PySimpleGUI as sg
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落对齐样式
from docx.enum.table import WD_ALIGN_VERTICAL  # 表格垂直对齐方式
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt  # 段落字体磅数
from docx.shared import RGBColor  # 字体颜色
from docx.shared import Inches  # 图标尺寸
from docx.shared import Cm  # 表格高度、宽度等(厘米 )
from docx.oxml.ns import qn  # 设置中文格式

start_t = datetime.datetime.now()

# 数据表格命名规则,excel文件类型为xls
'''
未入网、未上线数据——总表1.xls
轨迹漂移率、轨迹完成率数据——总表2.xls
数据合格率低于99.99%的车辆数据——总表3.xls
报警数据——总表4.xls
报警数据 行政审批局车辆--行政审批局车辆.xls
营运车辆查询-->车辆动态综合情况统-->车辆错误数据统计-->车辆报警信息统计
'''


path = 'D://专项整治//数据表//总表1.xls'
path2 = 'D://专项整治//数据表//总表2.xls'
path3 = 'D://专项整治//数据表//总表3.xls'
path4 = 'D://专项整治//数据表//总表4.xls'
path5 = 'D://专项整治//数据表//行政审批局车辆.xls'

# 打开总表1
workbook = xlrd.open_workbook(path)
sheet_name = workbook.sheet_names()
sheet = workbook.sheet_by_name(sheet_name[0])
# 打开总表2
workbook2 = xlrd.open_workbook(path2)
sheet_name2 = workbook2.sheet_names()
sheet2 = workbook2.sheet_by_name(sheet_name2[0])
# 打开总表3
workbook3 = xlrd.open_workbook(path3)
sheet_name3 = workbook3.sheet_names()
sheet3 = workbook3.sheet_by_name(sheet_name3[0])
# 打开总表4
workbook4 = xlrd.open_workbook(path4)
sheet_name4 = workbook4.sheet_names()
sheet4 = workbook4.sheet_by_name(sheet_name4[0])
# 打开总表5
workbook5 = xlrd.open_workbook(path5)
sheet_name5 = workbook5.sheet_names()
sheet5 = workbook5.sheet_by_name(sheet_name5[0])

print("------------------")
print("sheet页：" + str(sheet.name) + "\r\n文件行数：" + str(sheet.nrows) + " 行\r\n文件列数：" + str(sheet.ncols) + " 列")
print("------------------")
print("sheet页：" + str(sheet2.name) + "\r\n文件行数：" + str(sheet2.nrows) + " 行\r\n文件列数：" + str(sheet2.ncols) + " 列")
print("------------------")
print("sheet页：" + str(sheet3.name) + "\r\n文件行数：" + str(sheet3.nrows) + " 行\r\n文件列数：" + str(sheet3.ncols) + " 列")
print("------------------")
print("sheet页：" + str(sheet4.name) + "\r\n文件行数：" + str(sheet4.nrows) + " 行\r\n文件列数：" + str(sheet4.ncols) + " 列")
print("------------------")
print("sheet页：" + str(sheet5.name) + "\r\n文件行数：" + str(sheet5.nrows) + " 行\r\n文件列数：" + str(sheet5.ncols) + " 列")
print("------------------")



def createFileOfDirectory(n: str, m: str, now):
    '''
    创建生成文件的文件,如果父文件夹(n: str)存在，先删除父文件夹，备份数据文件
    :param n: 文件夹名
    :param m: 文件夹名
    :return:
    '''

    for i in city:
        path = "D://专项整治//{}//{}//{}".format(n, m, i.replace('市', ''))
        dir_parh = "D://专项整治//{}".format(n)
        if os.path.exists(path):
            shutil.rmtree(dir_parh)
        if not os.path.exists(path):
            os.makedirs(path)

    data_dir_path = "D://专项整治//数据表//"
    dest_data_dir_path = "D://专项整治//{}//".format(
        '数据表-' + now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日')
    if os.path.exists(dest_data_dir_path):
        shutil.rmtree(dest_data_dir_path)
    shutil.copytree(data_dir_path, dest_data_dir_path)

    return 0


NoOnlineExcel_dict = {}


def createFileOfNoOnlineExcel(excelName, fileName, fileName2, date_time):
    '''
    未上线车辆明细
    :param excelName: excel文件名
    :param fileName:
    :param fileName2:
    :return:
    '''
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '营口', excelName, '营口')
            count = select_db(
                "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(date_time,
                    '营口', '营口市运管处'))
            count2 = select_db(
                "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%包车客运%'".format(date_time,
                    '营口', '营口市运管处'))
            count3 = select_db(
                "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%危货运输%'".format(date_time,
                    '营口', '营口市运管处'))
            NoOnlineExcel_dict['营口市'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('营口市',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(date_time,
                        '营口', '营口市运管处'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%包车客运%'".format(date_time,
                        '营口', '营口市运管处'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%危货运输%'".format(date_time,
                        '营口', '营口市运管处'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        if cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '行政审批局', excelName, '行政审批局')
            count = select_db(
                "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(date_time,
                    '营口', '行政审批局'))
            count2 = select_db(
                "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%包车客运%'".format(date_time,
                    '营口', '行政审批局'))
            count3 = select_db(
                "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%危货运输%'".format(date_time,
                    '营口', '行政审批局'))
            NoOnlineExcel_dict['行政审批局'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('行政审批局',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(date_time,
                        '营口', '行政审批局'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%包车客运%'".format(date_time,
                        '营口', '行政审批局'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%危货运输%'".format(date_time,
                        '营口', '行政审批局'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                         cname.replace('市', ''))
        count = select_db(
            "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and industry_name like '%班车客运%'".format(date_time,
                cname))
        count2 = select_db(
            "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and industry_name like '%包车客运%'".format(date_time,
                cname))
        count3 = select_db(
            "select count(*) from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and industry_name like '%危货运输%'".format(date_time,
                cname))
        NoOnlineExcel_dict[cname] = [count[0][0], count2[0][0], count3[0][0]]
        # print(cname,count,count2,count3)
        workbook = openpyxl.Workbook()
        if count[0][0] > 0:
            sh = workbook.active
            sh.title = '班车客运'
            firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                        '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh.append(firstRow)
            value = select_db(
                "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and industry_name like '%班车客运%'".format(date_time,
                    cname))
            for i in value:
                sh.append(i)
            workbook.save(path)
        if count2[0][0] > 0:
            sh2 = workbook.create_sheet('包车客运')
            firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                        '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh2.append(firstRow)
            value = select_db(
                "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and industry_name like '%包车客运%'".format(date_time,
                    cname))
            for i in value:
                sh2.append(i)
            workbook.save(path)
        if count3[0][0] > 0:
            sh3 = workbook.create_sheet('危货运输')
            firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                        '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh3.append(firstRow)
            value = select_db(
                "select * from Ve_Regu_Details where strftime('%s',Last_online_date)<strftime('%s',date('{}')) and region_name like '%{}%' and industry_name like '%危货运输%'".format(date_time,
                    cname))
            for i in value:
                sh3.append(i)
            workbook.save(path)
    return 0


TrajectoryIntegrityRate_dict = {}


def createFileOfTrajectoryIntegrityRate(excelName, fileName, fileName2):
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '营口', excelName, '营口')
            count = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' and  运管机构名称 = '营口市运管处' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                    cname))
            count2 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' and  运管机构名称 = '营口市运管处' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                    cname))
            count3 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' and  运管机构名称 = '营口市运管处' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                    cname))
            TrajectoryIntegrityRate_dict['营口市'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('营口市',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                            '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '班车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                            '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '包车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                            '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '危货运输' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        if cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '行政审批局', excelName, '行政审批局')
            count = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' and  运管机构名称 = '{}' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                    '营口', '行政审批局'))
            count2 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' and  运管机构名称 = '{}' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                    '营口', '行政审批局'))
            count3 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' and  运管机构名称 = '{}' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                    '营口', '行政审批局'))
            TrajectoryIntegrityRate_dict['行政审批局'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('行政审批局',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                            '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '班车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                            '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '包车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                            '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '危货运输' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                         cname.replace('市', ''))
        count = select_db(
            "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                cname))
        count2 = select_db(
            "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                cname))
        count3 = select_db(
            "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ".format(
                cname))
        TrajectoryIntegrityRate_dict[cname] = [count[0][0], count2[0][0], count3[0][0]]
        # print(cname,count,count2,count3)
        workbook = openpyxl.Workbook()
        if count[0][0] > 0:
            sh = workbook.active
            sh.title = '班车客运'
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                        '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh.append(firstRow)
            value = select_db(
                "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                    cname))
            for i in value:
                sh.append(i)
            workbook.save(path)
        if count2[0][0] > 0:
            sh2 = workbook.create_sheet('包车客运')
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                        '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh2.append(firstRow)
            value = select_db(
                "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                    cname))
            for i in value:
                sh2.append(i)
            workbook.save(path)
        if count3[0][0] > 0:
            sh3 = workbook.create_sheet('危货运输')
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '行驶里程', '完整里程',
                        '轨迹完整率', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh3.append(firstRow)
            value = select_db(
                "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.行驶里程_公里 行驶里程,s.完整里程_公里 完整里程,s.轨迹完整率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' and s.行驶里程_公里  <> 0 and s.轨迹完整率 < 99.9 ORDER BY s.轨迹完整率 desc;".format(
                    cname))
            for i in value:
                sh3.append(i)
            workbook.save(path)
    return 0


TrajectoryDriftRate_dict = {}


def createFileOfTrajectoryDriftRate(excelName, fileName, fileName2):
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '营口', excelName, '营口')
            count = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' and  运管机构名称 != '行政审批局'  and s.轨迹漂移率> 0.4 ".format(
                    cname))
            count2 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' and  运管机构名称 != '行政审批局'  and s.轨迹漂移率> 0.4 ".format(
                    cname))
            count3 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' and  运管机构名称 != '行政审批局'  and s.轨迹漂移率> 0.4 ".format(
                    cname))
            TrajectoryDriftRate_dict['营口市'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('营口市',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                            '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  s.运管机构名称 = '{}' and s.所属行业 = '班车客运' and s.轨迹漂移率> 0.4 ORDER BY s.轨迹漂移率 desc;".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                            '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '包车客运' and s.轨迹漂移率> 0.4 ORDER BY s.轨迹漂移率 desc;".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                            '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '危货运输' and s.轨迹漂移率> 0.4 ORDER BY s.轨迹漂移率 desc;".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        if cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '行政审批局', excelName, '行政审批局')
            count = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' and  运管机构名称 = '{}'  and s.轨迹漂移率 > 0.4 ".format(
                    '营口', '行政审批局'))
            count2 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' and  运管机构名称 = '{}' and s.轨迹漂移率 > 0.4 ".format(
                    '营口', '行政审批局'))
            count3 = select_db(
                "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' and  运管机构名称 = '{}'  and s.轨迹漂移率 > 0.4 ".format(
                    '营口', '行政审批局'))
            TrajectoryDriftRate_dict['行政审批局'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('行政审批局',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                            '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '班车客运' and  s.轨迹漂移率> 0.4 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                            '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '包车客运' and s.轨迹漂移率> 0.4 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                            '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and  运管机构名称 = '{}' and s.所属行业 = '危货运输' and s.轨迹漂移率> 0.4 ORDER BY s.轨迹完整率 desc;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                         cname.replace('市', ''))
        count = select_db(
            "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '班车客运' AND s.轨迹漂移率> 0.4".format(
                cname))
        count2 = select_db(
            "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '包车客运' AND s.轨迹漂移率> 0.4 ".format(
                cname))
        count3 = select_db(
            "SELECT count(*)  FROM Ve_Status_Details s  WHERE s.车籍地 like '%{}%' and s.所属行业 = '危货运输' AND s.轨迹漂移率> 0.4 ".format(
                cname))
        TrajectoryDriftRate_dict[cname] = [count[0][0], count2[0][0], count3[0][0]]
        # print(cname,count,count2,count3)
        workbook = openpyxl.Workbook()
        if count[0][0] > 0:
            sh = workbook.active
            sh.title = '班车客运'
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                        '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh.append(firstRow)
            value = select_db(
                "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率 FROM Ve_Status_Details s WHERE s.车籍地 LIKE '%{}%' AND s.所属行业='班车客运' AND s.轨迹漂移率> 0.4 ORDER BY s.轨迹漂移率 DESC;".format(
                    cname))
            for i in value:
                sh.append(i)
            workbook.save(path)
        if count2[0][0] > 0:
            sh2 = workbook.create_sheet('包车客运')
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                        '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh2.append(firstRow)
            value = select_db(
                "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率 FROM Ve_Status_Details s WHERE s.车籍地 LIKE '%{}%' AND s.所属行业='包车客运' AND s.轨迹漂移率> 0.4 ORDER BY s.轨迹漂移率 DESC;".format(
                    cname))
            for i in value:
                sh2.append(i)
            workbook.save(path)
        if count3[0][0] > 0:
            sh3 = workbook.create_sheet('危货运输')
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '运管机构名称', '所属行业', '所属企业', '所属平台', '车籍地', '营运状态', '正常位置数', '轨迹漂移点数',
                        '轨迹漂移率', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh3.append(firstRow)
            value = select_db(
                "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.运管机构名称,s.所属行业,所属企业,s.所属平台,s.车籍地,s.营运状态,s.正常位置数,s.轨迹漂移点数,s.轨迹漂移率 FROM Ve_Status_Details s WHERE s.车籍地 LIKE '%{}%' AND s.所属行业='危货运输' AND s.轨迹漂移率> 0.4 ORDER BY s.轨迹漂移率 DESC;".format(
                    cname))
            for i in value:
                sh3.append(i)
            workbook.save(path)
    return 0


NoIntoNetworkExcel_dict = {}


def createFileOfNoIntoNetworkExcel(excelName, fileName, fileName2):
    '''
    未入网车辆明细
    :param excelName: 未入网excel文件名
    :param fileName:
    :param fileName2:
    :return:
    '''
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '营口', excelName, '营口')
            count = select_db(
                "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(
                    '营口', '营口市运管处'))
            count2 = select_db(
                "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%包车客运%'".format(
                    '营口', '营口市运管处'))
            count3 = select_db(
                "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%危货运输%'".format(
                    '营口', '营口市运管处'))
            NoIntoNetworkExcel_dict['营口市'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('营口市',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%包车客运%'".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%危货运输%'".format(
                        '营口', '营口市运管处'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        if cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '行政审批局', excelName, '行政审批局')
            count = select_db(
                "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(
                    '营口', '行政审批局'))
            count2 = select_db(
                "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%包车客运%'".format(
                    '营口', '行政审批局'))
            count3 = select_db(
                "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%危货运输%'".format(
                    '营口', '行政审批局'))
            NoIntoNetworkExcel_dict['行政审批局'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('行政审批局',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and industry_name like '%班车客运%'".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%包车客运%'".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                            '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and institution_name = '{}' and  industry_name like '%危货运输%'".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                         cname.replace('市', ''))
        count = select_db(
            "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and industry_name like '%班车客运%'".format(
                cname))
        count2 = select_db(
            "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and industry_name like '%包车客运%'".format(
                cname))
        count3 = select_db(
            "select count(*) from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and industry_name like '%危货运输%'".format(
                cname))
        NoIntoNetworkExcel_dict[cname] = [count[0][0], count2[0][0], count3[0][0]]
        # print(cname,count,count2,count3)
        workbook = openpyxl.Workbook()
        if count[0][0] > 0:
            sh = workbook.active
            sh.title = '班车客运'
            firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                        '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh.append(firstRow)
            value = select_db(
                "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and industry_name like '%班车客运%'".format(
                    cname))
            for i in value:
                sh.append(i)
            workbook.save(path)
        if count2[0][0] > 0:
            sh2 = workbook.create_sheet('包车客运')
            firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                        '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh2.append(firstRow)
            value = select_db(
                "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and industry_name like '%包车客运%'".format(
                    cname))
            for i in value:
                sh2.append(i)
            workbook.save(path)
        if count3[0][0] > 0:
            sh3 = workbook.create_sheet('危货运输')
            firstRow = ['车牌号码', '车牌颜色', '车辆属地管辖机构名称', '所属行业', '所属地区', '车辆类型', '业户名称', '平台名称', '入网状态', '年审截止日期',
                        '第一次入网时间', '最后上线日期', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh3.append(firstRow)
            value = select_db(
                "select * from Ve_Regu_Details where Network_status = '未入网' and region_name like '%{}%' and industry_name like '%危货运输%'".format(
                    cname))
            for i in value:
                sh3.append(i)
            workbook.save(path)
    return 0


VeConformityDetails_dict = {}


def createFileOfVeConformityDetails(excelName, fileName, fileName2):
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '营口', excelName, '营口')
            count = select_db(
                "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 管辖机构名称 != '{}' and 所属行业 like '%班车客运%'  AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                    '营口', '行政审批局'))
            count2 = select_db(
                "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 管辖机构名称 != '{}' and 所属行业 like '%包车客运%'  AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                    '营口', '行政审批局'))
            count3 = select_db(
                "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 管辖机构名称 != '{}' and 所属行业 like '%危货运输%'  AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                    '营口', '行政审批局'))
            VeConformityDetails_dict['营口市'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('营口市',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                            '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 管辖机构名称 !='{}' AND 所属行业 LIKE '%班车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                            '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 管辖机构名称 !='{}' AND 所属行业 LIKE '%包车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                            '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 管辖机构名称 !='{}' AND 所属行业 LIKE '%危货运输%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        if cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, '行政审批局', excelName, '行政审批局')
            count = select_db(
                "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 管辖机构名称 = '{}' and 所属行业 like '%班车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                    '营口', '行政审批局'))
            count2 = select_db(
                "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 管辖机构名称 = '{}' and 所属行业 like '%包车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                    '营口', '行政审批局'))
            count3 = select_db(
                "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 管辖机构名称 = '{}' and 所属行业 like '%危货运输%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                    '营口', '行政审批局'))
            VeConformityDetails_dict['行政审批局'] = [count[0][0], count2[0][0], count3[0][0]]
            # print('行政审批局',count, count2, count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                            '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 管辖机构名称 ='{}' AND 所属行业 LIKE '%班车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                            '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 管辖机构名称 ='{}' AND 所属行业 LIKE '%包车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                            '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 管辖机构名称 ='{}' AND 所属行业 LIKE '%危货运输%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                        '营口', '行政审批局'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
            continue
        path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                         cname.replace('市', ''))
        count = select_db(
            "select count(*) from Ve_Conformity_Details where 所属地区 like '%{}%' and 所属行业 like '%班车客运%'  AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                cname))
        count2 = select_db(
            "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 所属行业 like '%包车客运%'  AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                cname))
        count3 = select_db(
            "select count(*) from Ve_Conformity_Details where  所属地区 like '%{}%' and 所属行业 like '%危货运输%'  AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999".format(
                cname))
        VeConformityDetails_dict[cname] = [count[0][0], count2[0][0], count3[0][0]]
        # print(cname,count,count2,count3)
        workbook = openpyxl.Workbook()
        if count[0][0] > 0:
            sh = workbook.active
            sh.title = '班车客运'
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                        '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh.append(firstRow)
            value = select_db(
                "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 所属行业 LIKE '%班车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                    cname))
            for i in value:
                sh.append(i)
            workbook.save(path)
        if count2[0][0] > 0:
            sh2 = workbook.create_sheet('包车客运')
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                        '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh2.append(firstRow)
            value = select_db(
                "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 所属行业 LIKE '%包车客运%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                    cname))
            for i in value:
                sh2.append(i)
            workbook.save(path)
        if count3[0][0] > 0:
            sh3 = workbook.create_sheet('危货运输')
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属业户', '所属平台', '所属地区', '管辖机构名称', '所属行业', '位置总数', '正常位置数', '数据合格率',
                        '错误位置数', '经纬度错误数', '时间错误数', '速度错误数', '方向错误数', '海拔错误数', '处理情况', '处理手段', '无法处理原因', '处理时限']
            sh3.append(firstRow)
            value = select_db(
                "SELECT 日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称,所属行业 ,位置总数,正常位置数,数据合格率,错误位置数,经纬度错误数,时间错误数,速度错误数,方向错误数,海拔错误数,处理情况,处理手段,无法处理原因,处理时限 FROM Ve_Conformity_Details WHERE 所属地区 LIKE '%{}%' AND 所属行业 LIKE '%危货运输%' AND CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5))< 0.9999;".format(
                    cname))
            for i in value:
                sh3.append(i)
            workbook.save(path)
    return 0


Ve_Tired_Details_dict = {}


def createFileOfVeTiredDetails(excelName, fileName, fileName2):
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                             cname.replace('市', ''))
            count = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count2 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count3 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  超速报警数 <> 0;".format(
                    cname))
            count4 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count5 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count6 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  超速报警数 <> 0;".format(
                    cname))
            Ve_Tired_Details_dict[cname] = [count[0][0], count2[0][0], count3[0][0], count4[0][0], count5[0][0],
                                            count6[0][0]]
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) AND s.所属行业='班车客运' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        cname))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) AND s.所属行业='包车客运' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        cname))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) AND s.所属行业='危货运输' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        cname))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
        elif cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                             cname.replace('市', ''))
            count = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count2 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count3 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  超速报警数 <> 0;".format(
                    cname))
            count4 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count5 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  超速报警数 <> 0;".format(
                    cname))
            count6 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  超速报警数 <> 0;".format(
                    cname))
            Ve_Tired_Details_dict[cname] = [count[0][0], count2[0][0], count3[0][0], count4[0][0], count5[0][0],
                                            count6[0][0]]
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段',
                            '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) AND s.所属行业='班车客运' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        '营口'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段',
                            '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) AND s.所属行业='包车客运' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        '营口'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段',
                            '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) AND s.所属行业='危货运输' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        '营口'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
        else:
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                             cname.replace('市', ''))
            count = select_db(
                "select count(*) from Ve_Tired_Details where 所属地区 like '%{}%' and 所属行业 like '%班车客运%' and  超速报警数 <> 0".format(
                    cname))
            count2 = select_db(
                "select count(*) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%包车客运%' and  超速报警数 <> 0".format(
                    cname))
            count3 = select_db(
                "select count(*) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%危货运输%' and  超速报警数 <> 0".format(
                    cname))
            count4 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%班车客运%' and  超速报警数 <> 0".format(
                    cname))
            count5 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%包车客运%' and  超速报警数 <> 0".format(
                    cname))
            count6 = select_db(
                "select ifnull(sum(超速报警数),0) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%危货运输%' and  超速报警数 <> 0".format(
                    cname))
            Ve_Tired_Details_dict[cname] = [count[0][0], count2[0][0], count3[0][0], count4[0][0], count5[0][0],
                                            count6[0][0]]
            # print(cname,count,count2,count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' AND s.所属行业='班车客运' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        cname))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' AND s.所属行业='包车客运' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        cname))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.超速报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' AND s.所属行业='危货运输' AND s.超速报警数<> 0 ORDER BY s.超速报警数 desc".format(
                        cname))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
    return 0


Ve_Tired_Details_dict2 = {}


def createFileOfVeTiredDetails2(excelName, fileName, fileName2):
    for cname in city:
        if cname == '营口市':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                             cname.replace('市', ''))
            count = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count2 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count3 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count4 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count5 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count6 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            Ve_Tired_Details_dict2[cname] = [count[0][0], count2[0][0], count3[0][0], count4[0][0], count5[0][0],
                                             count6[0][0]]
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) AND s.所属行业='班车客运' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        cname))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) AND s.所属行业='包车客运' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        cname))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码 not in (SELECT * FROM Ve_xingzheng) AND s.所属行业='危货运输' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        cname))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
        elif cname == '行政审批局':
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                             cname.replace('市', ''))
            count = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count2 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count3 = select_db(
                "select count(*) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count4 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%班车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count5 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%包车客运%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            count6 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details s where  所属地区 like '%营口%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) and 所属行业 like '%危货运输%' and  疲劳驾驶报警数 <> 0;".format(
                    cname))
            Ve_Tired_Details_dict2[cname] = [count[0][0], count2[0][0], count3[0][0], count4[0][0], count5[0][0],
                                             count6[0][0]]
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段',
                            '无法处理原因', '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) AND s.所属行业='班车客运' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        '营口'))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段',
                            '无法处理原因', '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) AND s.所属行业='包车客运' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        '营口'))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段',
                            '无法处理原因', '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' and 车牌号码  in (SELECT * FROM Ve_xingzheng) AND s.所属行业='危货运输' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        '营口'))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
        else:
            path = "D://专项整治//{}//{}//{}//{}-{}.xlsx".format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                             cname.replace('市', ''))
            count = select_db(
                "select count(*) from Ve_Tired_Details where 所属地区 like '%{}%' and 所属行业 like '%班车客运%' and  疲劳驾驶报警数 <> 0".format(
                    cname))
            count2 = select_db(
                "select count(*) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%包车客运%' and  疲劳驾驶报警数 <> 0".format(
                    cname))
            count3 = select_db(
                "select count(*) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%危货运输%' and  疲劳驾驶报警数 <> 0".format(
                    cname))
            count4 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%班车客运%' and  疲劳驾驶报警数 <> 0".format(
                    cname))
            count5 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%包车客运%' and  疲劳驾驶报警数 <> 0".format(
                    cname))
            count6 = select_db(
                "select ifnull(sum(疲劳驾驶报警数),0) from Ve_Tired_Details where  所属地区 like '%{}%' and 所属行业 like '%危货运输%' and  疲劳驾驶报警数 <> 0".format(
                    cname))
            Ve_Tired_Details_dict2[cname] = [count[0][0], count2[0][0], count3[0][0], count4[0][0], count5[0][0],
                                             count6[0][0]]
            # print(cname,count,count2,count3)
            workbook = openpyxl.Workbook()
            if count[0][0] > 0:
                sh = workbook.active
                sh.title = '班车客运'
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' AND s.所属行业='班车客运' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        cname))
                for i in value:
                    sh.append(i)
                workbook.save(path)
            if count2[0][0] > 0:
                sh2 = workbook.create_sheet('包车客运')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh2.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' AND s.所属行业='包车客运' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        cname))
                for i in value:
                    sh2.append(i)
                workbook.save(path)
            if count3[0][0] > 0:
                sh3 = workbook.create_sheet('危货运输')
                firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '疲劳驾驶报警数', '处理情况', '处理手段', '无法处理原因',
                            '处理时限']
                sh3.append(firstRow)
                value = select_db(
                    "SELECT s.日期范围,s.车牌号码,s.车牌颜色,s.所属行业,s.所属企业,s.所属平台,s.所属地区,s.疲劳驾驶报警数 FROM Ve_Tired_Details s WHERE s.所属地区 like '%{}%' AND s.所属行业='危货运输' AND s.疲劳驾驶报警数<> 0 ORDER BY s.疲劳驾驶报警数 desc".format(
                        cname))
                for i in value:
                    sh3.append(i)
                workbook.save(path)
    return 0


def select_db(n):
    # return type [(,),]
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = n
    cursor.execute(sql)
    values = cursor.fetchall()
    cursor.close()
    pythondb.close()
    return values


def truncate_db(n):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "DELETE FROM {};".format(n)
    cursor.execute(sql)
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return 0


def count_db(n):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "select count(*) from  {};".format(n)
    cursor.execute(sql)
    count = cursor.fetchall()
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return count[0][0]


def insert_db(n, m):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "insert into {}(Licenseplate_numbe, Licenseplate_colo, institution_name, industry_name, region_name, type_name, Business_name, Platform_name, Network_status, Annualreview_exdate, enternetwork_fdate, Last_online_date, dispose, means, Unable_reason, time_limit) values (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) ".format(
        m)
    cursor.executemany(sql, n)
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return 0


def insert_db2(n, m):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "insert into {}(日期范围, 车牌号码, 车牌颜色, 运管机构名称, 所属行业, 所属企业, 所属平台, 车籍地, 营运状态, 照片数, 车辆上线天数, 在线时长_分钟, 行驶里程_公里, 位置总数, 正常位置数, 错误位置数, 数据合格率, 轨迹漂移点数, 不完整里程_公里, 完整里程_公里, 轨迹完整率, 里程是否完整, 轨迹漂移率, 异常点数, 超速总次数, 疲劳报警时长_分钟, 疲劳报警次数, 报警总数, 轨迹合格率, 处理情况, 处理手段, 无法处理原因, 处理时限) values (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) ".format(
        m)
    cursor.executemany(sql, n)
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return 0


def insert_db3(n, m):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "insert into {}(日期范围, 车牌号码, 车牌颜色, 所属业户, 所属平台, 所属地区, 管辖机构名称, 所属行业, 位置总数, 正常位置数, 错误位置数, 经纬度错误数, 时间错误数, 速度错误数, 方向错误数, 海拔错误数, 处理情况, 处理手段, 无法处理原因, 处理时限) values (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) ".format(
        m)
    cursor.executemany(sql, n)
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return 0


def insert_db4(n, m):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "insert into {}(日期范围, 车牌号码, 车牌颜色, 所属行业, 所属企业, 所属平台, 所属地区, 报警总数, 紧急报警数, 超速报警数, 疲劳驾驶报警数, 夜间行驶报警数, 禁入报警数, 禁出报警数, 偏航报警数, 超范围经营报警数,高速公路违停报警,定位数据异常报警, 处理情况, 处理手段, 无法处理原因, 处理时限) values (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) ".format(
        m)
    cursor.executemany(sql, n)
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return 0


def insert_db5(n, m):
    pythondb = sqlite3.connect("D://python//Ve_Regu_Details.db")
    cursor = pythondb.cursor()
    sql = "insert into {}(车牌号码) values (?) ".format(m)
    cursor.executemany(sql, n)
    pythondb.commit()
    cursor.close()
    pythondb.close()
    return 0


def VeConformityDetails_percent(excelName, fileName, fileName2):
    for cname in city:
        path = 'D://专项整治//{}//{}//{}//{}-{}.xlsx'.format(fileName, fileName2, cname.replace('市', ''), excelName,
                                                         cname.replace('市', ''))

        # 打开数据合格率低于99.99%的车辆明细表
        workbook = openpyxl.load_workbook(path)
        sheet_name = workbook.sheetnames

        for i in range(3):
            try:
                sheet = workbook[sheet_name[i]]
                row_number = sheet.max_row
                if row_number == 1 or row_number == 0:
                    continue
                for k in range(2, row_number + 1):
                    # "%.2f%%" % (JI * 100)
                    JI = "%.2f%%" % ((sheet['J' + str(k)].value / sheet['I' + str(k)].value) * 100)
                    if JI == '99.99%':
                        JI = '99.98%'
                    call = 'K' + str(k)
                    # formula = '=J' + str(k) + '/I' + str(k)
                    sheet[call] = JI
                workbook.save(path)
            except:
                continue


def alarmOverspeedTired_yingkou_xingzhengshenpiju(excelName, fileName, fileName2):
    '''
    超速和疲劳报警中，生成行政审批局的数据excel
    '''
    path = 'D://专项整治//数据表//行政审批局车牌号.xls'

    # openpyxl不能打开xls格式文件
    # workbook = openpyxl.load_workbook(path)
    # sheet = workbook[workbook.sheetnames[0]]

    workbook = xlrd.open_workbook(path)
    sheet_name = workbook.sheet_names()
    sheet = workbook.sheet_by_name(sheet_name[0])

    car_number = []
    for i in range(1, sheet.nrows):
        car_number.append(sheet.cell(i, 0).value)
    print(car_number)
    path2 = 'D://专项整治//{}//{}//营口//{}-营口.xlsx'.format(fileName, fileName2, excelName)
    path3 = 'D://专项整治//{}//{}//行政审批局//{}-行政审批局.xlsx'.format(fileName, fileName2, excelName)

    workbook = xlrd.open_workbook(path2)
    sheet_name = workbook.sheet_names()

    workbook2 = openpyxl.Workbook()
    for k, i in enumerate(sheet_name):
        sheet = workbook.sheet_by_name(i)
        if sheet.nrows == 0 or sheet.nrows == 1:
            continue
        if k == 0:
            sh = workbook2.active
            sh = workbook2.create_sheet(i, 0)
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                        '处理时限']
            sh.append(firstRow)

            workbook_xlrd = xlrd.open_workbook(path2)
            # sheet_name = workbook.sheet_names()
            sheet_xlrd = workbook_xlrd.sheet_by_name(i)

            for data_row in range(1, sheet_xlrd.nrows):
                print(1, sheet.cell(data_row, 1).value)
                if car_number.count(sheet.cell(data_row, 1).value) >= 1:
                    print('--', sheet.cell(data_row, 1).value)
                    hang = [str(sheet_xlrd.cell_value(data_row, i)).replace('.0', '') for i in
                            range(0, sheet_xlrd.ncols)]
                    sh.append(hang)

                    workbook3 = openpyxl.load_workbook(path2)
                    sheet_name = workbook3.sheetnames
                    sheet1 = workbook3[workbook3.sheetnames[0]]
                    # sheet1.delete_rows(data_row + 1, 1)
                    workbook3.save(path2)

            workbook2.save(path3)
        else:
            sh = workbook2.create_sheet(i, k)
            firstRow = ['日期范围', '车牌号码', '车牌颜色', '所属行业', '所属企业', '所属平台', '所属地区', '超速报警数', '处理情况', '处理手段', '无法处理原因',
                        '处理时限']
            sh.append(firstRow)

            workbook_xlrd = xlrd.open_workbook(path2)
            # sheet_name = workbook.sheet_names()
            sheet_xlrd = workbook_xlrd.sheet_by_name(i)

            for data_row in range(1, sheet_xlrd.nrows):
                print(2, sheet.cell(data_row, 1).value)
                if car_number.count(sheet.cell(data_row, 1).value) >= 1:
                    # print(sheet.cell(data_row, 1).value)
                    hang = [str(sheet_xlrd.cell_value(data_row, i)).replace('.0', '') for i in
                            range(0, sheet_xlrd.ncols)]
                    sh.append(hang)

                    workbook3 = openpyxl.load_workbook(path2)
                    sheet_name = workbook3.sheetnames
                    sheet1 = workbook3[workbook3.sheetnames[0]]
                    sheet1.delete_rows(data_row + 1, 1)
                    workbook3.save(path2)
            print('---')

            workbook2.save(path3)
    else:
        # 删除没有数据的sheet页
        workbook2 = openpyxl.load_workbook(path3)
        sheet_name = workbook2.sheetnames
        for i in sheet_name:
            sheet = workbook2[i]
            if sheet.max_row == 1 or sheet.max_row == 0:
                sheet = workbook2[i]
                workbook2.remove(sheet)
        workbook2.save(path3)

def get_current_week(n):
    '''

    :param n: 星期几(0-6)
    :return: 返回当前周对应星期的日期
    '''

    today_date = datetime.date.today()
    one_day = datetime.timedelta(days=1)

    if n == 0:
        while today_date.weekday() != 0:
            today_date -= one_day
        week_date = today_date
    elif n == 1:
        while today_date.weekday() != 1:
            today_date -= one_day
        week_date = today_date
    elif n == 2:
        while today_date.weekday() != 2:
            today_date -= one_day
        week_date = today_date
    elif n == 3:
        while today_date.weekday() != 3:
            today_date -= one_day
        week_date = today_date
    elif n == 4:
        while today_date.weekday() != 4:
            today_date -= one_day
        week_date = today_date
    elif n == 5:
        while today_date.weekday() != 5:
            today_date -= one_day
        week_date = today_date
    elif n == 6:
        while today_date.weekday() != 6:
            today_date += one_day
        week_date = today_date

    return week_date


# 未入网、未上线数据入sqlite
truncate_db('Ve_Regu_Details')
i = 1
while True:
    a = []
    if i + 2000 <= sheet.nrows:
        for i in range(i, i + 2000):
            rows = sheet.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db(a, 'Ve_Regu_Details')
    else:
        for i in range(i, sheet.nrows):
            rows = sheet.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db(a, 'Ve_Regu_Details')
        break
    i += 1
num = count_db('Ve_Regu_Details')
print("总共执行了 {} 条未入网、未上线数据".format(num))
# 轨迹漂移率、轨迹完成率数据入sqlite
truncate_db('Ve_Status_Details')
i = 1
while True:
    a = []
    if i + 2000 <= sheet2.nrows:
        for i in range(i, i + 2000):
            rows = sheet2.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db2(a, 'Ve_Status_Details')
    else:
        for i in range(i, sheet2.nrows):
            rows = sheet2.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db2(a, 'Ve_Status_Details')
        break
    i += 1
num = count_db('Ve_Status_Details')
print("总共执行了 {} 条轨迹漂移率、轨迹完成率数据".format(num))
# 数据合格率低于99.99%的车辆数据入sqlite
truncate_db('Ve_Conformity_Details')
i = 1
while True:
    a = []
    if i + 2000 <= sheet3.nrows:
        for i in range(i, i + 2000):
            rows = sheet3.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db3(a, 'Ve_Conformity_Details')
    else:
        for i in range(i, sheet3.nrows):
            rows = sheet3.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db3(a, 'Ve_Conformity_Details')
        break
    i += 1
num = count_db('Ve_Conformity_Details')
print("总共执行了 {} 条数据合格率低于99.99%的车辆数据".format(num))
# 报警数据入sqlite
truncate_db('Ve_Tired_Details')
i = 1
while True:
    a = []
    if i + 2000 <= sheet4.nrows:
        for i in range(i, i + 2000):
            rows = sheet4.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db4(a, 'Ve_Tired_Details')
    else:
        for i in range(i, sheet4.nrows):
            rows = sheet4.row_values(i)
            [rows.append('') for i in range(4)]
            a.append(rows)
        insert_db4(a, 'Ve_Tired_Details')
        break
    i += 1
num = count_db('Ve_Tired_Details')
print("总共执行了 {} 条报警数据".format(num))
# 报警数据入sqlite
truncate_db('Ve_XingZheng')
i = 1
while True:
    a = []
    if i + 2000 <= sheet5.nrows:
        for i in range(i, i + 2000):
            rows = [sheet5.cell(i, 0).value]
            a.append(rows)
        insert_db5(a, 'Ve_XingZheng')
    else:
        for i in range(i, sheet5.nrows):
            rows = [sheet5.cell(i, 0).value]
            a.append(rows)
        insert_db5(a, 'Ve_XingZheng')
        break
    i += 1
num = count_db('Ve_XingZheng')
print("总共执行了 {} 条行政审批局车牌号数据".format(num))

print("------------------")

# 生成文件的路径文件夹名称
# now = datetime.datetime.now()
# now = (datetime.datetime.now() - datetime.timedelta(days=1))
# oneWeekDayAgo = (datetime.datetime.now() - datetime.timedelta(days=7))

# now = 上一周的星期日的日期，oneWeekDayAgo = 上一周星期一的日期
now = get_current_week(0)-datetime.timedelta(days=1)
oneWeekDayAgo = (get_current_week(0) - datetime.timedelta(days=7))

fileName = "专项整治{}-{}".format(
    oneWeekDayAgo.strftime('%Y') + '年' + oneWeekDayAgo.strftime('%m') + '月' + oneWeekDayAgo.strftime('%d') + '日',
    now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日')
fileName2 = "2-各市车辆问题明细{}-{}".format(
    oneWeekDayAgo.strftime('%Y') + '年' + oneWeekDayAgo.strftime('%m') + '月' + oneWeekDayAgo.strftime('%d') + '日',
    now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日')

# 设置地址名称列表
city = ['沈阳市', '大连市', '鞍山市', '抚顺市', '本溪市', '丹东市', '锦州市', '营口市', '阜新市', '辽阳市', '铁岭市', '朝阳市', '盘锦市', '葫芦岛市', '行政审批局']


# 创建文件生成新目标文件夹，备份数据文件夹到新文件夹
createFileOfDirectory(fileName, fileName2,now)



# 生成未上线车辆明细excel
createFileOfNoOnlineExcel('未上线车辆明细', fileName, fileName2,date_time=oneWeekDayAgo)
print('生成"未上线车辆明细.xlsx"完毕')

time.sleep(1)

# 生成未入网车辆明细excel
createFileOfNoIntoNetworkExcel('未入网车辆明细', fileName, fileName2)
print('生成"未入网车辆明细.xlsx"完毕\t', end="")
[print(i + '存在未入网车辆\t', end="") for i in city if
 NoIntoNetworkExcel_dict[i][0] > 0 or NoIntoNetworkExcel_dict[i][1] > 0 or NoIntoNetworkExcel_dict[i][2] > 0]
print("")

time.sleep(1)

# 生成轨迹完成率明细excel
createFileOfTrajectoryIntegrityRate('轨迹完成率明细', fileName, fileName2)
print('生成"轨迹完成率明细.xlsx"完毕')

time.sleep(1)

# 生成轨迹漂移率明细excel
createFileOfTrajectoryDriftRate('轨迹漂移率明细', fileName, fileName2)
print('生成"轨迹漂移率明细.xlsx"完毕')

time.sleep(1)

# 数据合格率低于99.99%的车辆明细excel
createFileOfVeConformityDetails('数据合格率低于99.99%的车辆明细', fileName, fileName2)
time.sleep(0.5)
# 数据合格率低于99.99%的车辆明细K列添加百分比数值
VeConformityDetails_percent('数据合格率低于99.99%的车辆明细', fileName, fileName2)
print('生成"数据合格率低于99.99%的车辆明细.xlsx"完毕')

time.sleep(1)

# 超速报警明细excel
createFileOfVeTiredDetails('超速报警明细', fileName, fileName2)
# alarmOverspeedTired_yingkou_xingzhengshenpiju('超速报警明细', fileName, fileName2)
print('生成"超速报警明细.xlsx"完毕')

time.sleep(1)

# 疲劳报警明细excel
createFileOfVeTiredDetails2('疲劳报警明细', fileName, fileName2)
# alarmOverspeedTired_yingkou_xingzhengshenpiju('疲劳报警明细', fileName, fileName2)
print('生成"疲劳报警明细.xlsx"完毕')

end_t = datetime.datetime.now()
finish_time = str((end_t - start_t).seconds)

start_t = datetime.datetime.now()


# 返回地市排名
def getSortByDict(n, m):
    '''
    返回字典排名
    :param n: 字典名
    :param m: 地市名
    :return:
    '''
    dictName = n
    cityName = m
    sort_dict = {}
    sort_list = []
    for i in dictName:
        sum = dictName[i][0] + dictName[i][1] + dictName[i][2]
        sort_list.append(sum)
        sort_dict[i] = sum
    sort_list.sort()
    return sort_list.index(sort_dict[cityName]) + 1


# 问题清单.docx
def createWord(n, m, fileName, fileName2):
    # 日期、时间相关变量
    # now=datetime.datetime.now()
    now = n
    oneWeekDayAgo = m
    tadayTime = now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日'
    tadaySevenAgoTime = oneWeekDayAgo.strftime('%Y') + '年' + oneWeekDayAgo.strftime(
        '%m') + '月' + oneWeekDayAgo.strftime('%d') + '日'

    # path='D://专项整治//问题清单.docx'
    path = 'D://专项整治//{}//{}//问题清单.docx'.format(fileName, fileName2)

    # 打开一个doc文件
    doc = Document()
    # 设置doc默认字体为宋体
    # 西文字体
    doc.styles['Normal'].font.name = u'宋体'
    # 中文字体
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 设置页边距
    sections = doc.sections[0]
    sections.left_margin = Cm(1.9)
    sections.right_margin = Cm(1.9)

    paragraph = doc.add_paragraph()
    # 设置标题居中，不加默认左对齐

    ##################全省车辆入网率问题清单############################################################
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆入网率问题清单\r')
    run.font.name = '方正小标宋简体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)
    # paragraph = doc.add_paragraph()
    # paragraph.add_run('统计周期：{}-{}'.format(tadaySevenAgoTime,tadayTime).font.size = Pt(12)

    # 添加word表格(Table Grid样式表格)
    table = doc.add_table(rows=19, cols=5, style='Table Grid')
    # 设置行高
    for i in range(19):
        table.rows[i].height = Cm(1)


    # 合并单元格
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(18, 1).merge(table.cell(18, 4))
    # 在单元格中插入内容
    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    # 单元格内容水平居中
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 单元格内容垂直居中
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER   #TOP、CENTER和BOTTOM
    # 字体格式
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    # 单元格内容水平居中
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 单元格内容垂直居中
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 字体格式
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(0, 4).paragraphs[0].add_run('合计（辆）')
    # 单元格内容水平居中
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 单元格内容垂直居中
    table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 字体格式
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    content = ['班车(辆)', '包车(辆)', '危货(辆)']
    for k, v in enumerate(content):
        # 在单元格中插入内容
        table_run = table.cell(1, k + 1).paragraphs[0].add_run(v)
        # 单元格内容水平居中
        table.cell(1, k + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 单元格内容垂直居中
        table.cell(1, k + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 字体格式
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    city = ['沈阳市', '大连市', '鞍山市', '抚顺市', '本溪市', '丹东市', '锦州市', '营口市', '阜新市', '辽阳市', '铁岭市', '朝阳市', '盘锦市', '葫芦岛市', '行政审批局']
    # 地址名称写入表
    for k, v in enumerate(city):
        # 在单元格中插入内容
        table_run = table.cell(k + 2, 0).paragraphs[0].add_run(v)
        # 单元格内容水平居中
        table.cell(k + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 单元格内容垂直居中
        table.cell(k + 2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 字体格式
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(17, 0).paragraphs[0].add_run('小计(辆)')
    # 单元格内容水平居中
    table.cell(17, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 单元格内容垂直居中
    table.cell(17, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 字体格式
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(18, 0).paragraphs[0].add_run('总计(辆)')
    # 单元格内容水平居中
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 单元格内容垂直居中
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 字体格式
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 各地市班车（辆）、包车（辆）、危货（辆）、合计（辆）
    for k, v in enumerate(city):
        for i in range(4):
            if i != 3:
                # 在单元格中插入内容
                # table_run = table.cell(k + 2, i + 1).paragraphs[0].add_run(str(NoIntoNetworkExcel_dict[v][i]).replace('0', '-'))
                table_run = table.cell(k + 2, i + 1).paragraphs[0].add_run(str('-' if NoIntoNetworkExcel_dict[v][i] == 0 else NoIntoNetworkExcel_dict[v][i]))
                # 单元格内容水平居中
                table.cell(k + 2, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 单元格内容垂直居中
                table.cell(k + 2, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 字体格式
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)
            else:
                # 在单元格中插入内容
                number = NoIntoNetworkExcel_dict[v][0] + NoIntoNetworkExcel_dict[v][1] + NoIntoNetworkExcel_dict[v][2]
                table_run = table.cell(k + 2, 4).paragraphs[0].add_run(str('-' if number == 0 else number))
                # 单元格内容水平居中
                table.cell(k + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 单元格内容垂直居中
                table.cell(k + 2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 字体格式
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)

    # 小计（辆）
    sum = 0
    sum = 0
    for i in range(4):
        sum = 0
        if i != 3:
            num = [NoIntoNetworkExcel_dict[j][i] for j in NoIntoNetworkExcel_dict]
            for j in num:
                sum += j
            table_run = table.cell(17, i + 1).paragraphs[0].add_run(str(sum))
            # 单元格内容水平居中
            table.cell(17, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 单元格内容垂直居中
            table.cell(17, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 字体格式
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
        else:
            for j in NoIntoNetworkExcel_dict:
                for k in range(3):
                    sum += NoIntoNetworkExcel_dict[j][k]
            table_run = table.cell(18, 1).paragraphs[0].add_run(str(sum))
            # 单元格内容水平居中
            table.cell(18, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 单元格内容垂直居中
            table.cell(18, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 字体格式
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
    else:
        table_run = table.cell(17, 4).paragraphs[0].add_run('-')
        # 单元格内容水平居中
        table.cell(17, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 单元格内容垂直居中
        table.cell(17, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 字体格式
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    ##################全省车辆上线率问题清单############################################################
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆上线率问题清单\r')
    run.font.name = '方正小标宋简体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)

    # 添加word表格(Table Grid央视表格)
    table = doc.add_table(rows=19, cols=6, style='Table Grid')
    # 设置行高
    for i in range(19):
        table.rows[i].height = Cm(1)

    # 合并单元格
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(18, 1).merge(table.cell(18, 5))

    # 在单元格中插入内容
    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    # 单元格内容水平居中
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 单元格内容垂直居中
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 字体格式
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 4).paragraphs[0].add_run('合计（辆）')
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 5).paragraphs[0].add_run('排名')
    table.cell(0, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 0).paragraphs[0].add_run('小计(辆)')
    table.cell(17, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 0).paragraphs[0].add_run('总计(辆)')
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 5).paragraphs[0].add_run('-')
    table.cell(17, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    content = ['班车(辆)', '包车(辆)', '危货(辆)']
    for k, v in enumerate(content):
        table_run = table.cell(1, k + 1).paragraphs[0].add_run(v)
        table.cell(1, k + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, k + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    city = ['沈阳市', '大连市', '鞍山市', '抚顺市', '本溪市', '丹东市', '锦州市', '营口市', '阜新市', '辽阳市', '铁岭市', '朝阳市', '盘锦市', '葫芦岛市', '行政审批局']
    for k, v in enumerate(city):
        table_run = table.cell(k + 2, 0).paragraphs[0].add_run(v)
        table.cell(k + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    # 各地市班车（辆）、包车（辆）、危货（辆）、合计（辆）
    for k, v in enumerate(city):
        for i in range(4):
            if i != 3:
                number = '-' if NoOnlineExcel_dict[v][i] == 0 else NoOnlineExcel_dict[v][i]
                table_run = table.cell(k + 2, i + 1).paragraphs[0].add_run(str(number))
                table.cell(k + 2, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)
            else:
                number = '-' if NoOnlineExcel_dict[v][0] + NoOnlineExcel_dict[v][1] + NoOnlineExcel_dict[v][2] == 0 else \
                    NoOnlineExcel_dict[v][0] + NoOnlineExcel_dict[v][1] + NoOnlineExcel_dict[v][2]
                table_run = table.cell(k + 2, 4).paragraphs[0].add_run(str(number))
                table.cell(k + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)

    # 排名
    for k, v in enumerate(city):
        number = getSortByDict(NoOnlineExcel_dict, v)
        table_run = table.cell(k + 2, 5).paragraphs[0].add_run(str(number))
        table.cell(k + 2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 小计（辆）
    sum = 0
    for i in range(4):
        sum = 0
        if i != 3:
            num = [NoOnlineExcel_dict[j][i] for j in NoOnlineExcel_dict]
            for j in num:
                sum += j
            table_run = table.cell(17, i + 1).paragraphs[0].add_run(str(sum))
            table.cell(17, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(17, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
        else:
            for j in NoOnlineExcel_dict:
                for k in range(3):
                    sum += NoOnlineExcel_dict[j][k]
            table_run = table.cell(18, 1).paragraphs[0].add_run(str(sum))
            table.cell(18, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
    else:
        table_run = table.cell(17, 4).paragraphs[0].add_run('-')
        table.cell(17, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(17, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)
    ##################全省车辆数据合格率问题清单############################################################
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆数据合格率问题清单')
    run.font.name = '方正小标宋简体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)

    # 添加word表格(Table Grid央视表格)
    table = doc.add_table(rows=20, cols=6, style='Table Grid')
    # 设置行高
    for i in range(20):
        table.rows[i].height = Cm(0.98)

    # 合并单元格
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(18, 1).merge(table.cell(18, 5))
    table.cell(19, 0).merge(table.cell(19, 5))

    # 在单元格中插入内容
    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 4).paragraphs[0].add_run('合计（辆）')
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 5).paragraphs[0].add_run('排名')
    table.cell(0, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 0).paragraphs[0].add_run('小计(辆)')
    table.cell(17, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 0).paragraphs[0].add_run('总计(辆)')
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 5).paragraphs[0].add_run('-')
    table.cell(17, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    content = ['班车(辆)', '包车(辆)', '危货(辆)']
    for k, v in enumerate(content):
        table_run = table.cell(1, k + 1).paragraphs[0].add_run(v)
        table.cell(1, k + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, k + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    for k, v in enumerate(city):
        table_run = table.cell(k + 2, 0).paragraphs[0].add_run(v)
        table.cell(k + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    # 各地市班车（辆）、包车（辆）、危货（辆）、合计（辆）
    for k, v in enumerate(city):
        for i in range(4):
            if i != 3:
                number = '-' if VeConformityDetails_dict[v][i] == 0 else VeConformityDetails_dict[v][i]
                table_run = table.cell(k + 2, i + 1).paragraphs[0].add_run(str(number))
                table.cell(k + 2, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)
            else:
                number = '-' if VeConformityDetails_dict[v][0] + VeConformityDetails_dict[v][1] + \
                                VeConformityDetails_dict[v][2] == 0 else VeConformityDetails_dict[v][0] + \
                                                                         VeConformityDetails_dict[v][1] + \
                                                                         VeConformityDetails_dict[v][2]
                table_run = table.cell(k + 2, 4).paragraphs[0].add_run(str(number))
                table.cell(k + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)
    # 小计（辆）
    sum = 0
    for i in range(4):
        sum = 0
        if i != 3:
            num = [VeConformityDetails_dict[j][i] for j in VeConformityDetails_dict]
            for j in num:
                sum += j
            table_run = table.cell(17, i + 1).paragraphs[0].add_run(str(sum))
            table.cell(17, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(17, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
        else:
            for j in VeConformityDetails_dict:
                for k in range(3):
                    sum += VeConformityDetails_dict[j][k]
            table_run = table.cell(18, 1).paragraphs[0].add_run(str(sum))
            table.cell(18, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
            # 备注数量
            ## 低于99.95车辆数量
            count = select_db(
                'SELECT count(*) FROM Ve_Conformity_Details WHERE CAST (CAST (正常位置数 AS float)/位置总数 AS decimal (10,5)) < 0.9995')
            table_run = table.cell(19, 0).paragraphs[0].add_run(
                '备注：数据合格率低于99.99%的车辆总数{}台，低于99.95车辆总数{}台。'.format(sum, count[0][0]))
            table.cell(19, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(19, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = '仿宋_GB2312'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            table_run.font.size = Pt(12)
    else:
        table_run = table.cell(17, 4).paragraphs[0].add_run('-')
        table.cell(17, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(17, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 排名
    for k, v in enumerate(city):
        number = getSortByDict(VeConformityDetails_dict, v)
        table_run = table.cell(k + 2, 5).paragraphs[0].add_run(str(number))
        table.cell(k + 2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)
    #################全省车辆轨迹完整率问题清单 #######################################################################
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆轨迹完整率问题清单')
    run.font.name = '方正小标宋简体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)

    # 添加word表格(Table Grid央视表格)
    table = doc.add_table(rows=20, cols=6, style='Table Grid')
    # 设置行高
    for i in range(20):
        table.rows[i].height = Cm(0.98)

    # 合并单元格
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(18, 1).merge(table.cell(18, 5))
    table.cell(19, 0).merge(table.cell(19, 5))

    # 在单元格中插入内容
    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 4).paragraphs[0].add_run('合计（辆）')
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 5).paragraphs[0].add_run('排名')
    table.cell(0, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 0).paragraphs[0].add_run('小计(辆)')
    table.cell(17, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 0).paragraphs[0].add_run('总计(辆)')
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 4).paragraphs[0].add_run('-')
    table.cell(17, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 5).paragraphs[0].add_run('-')
    table.cell(17, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    content = ['班车(辆)', '包车(辆)', '危货(辆)']
    for k, v in enumerate(content):
        table_run = table.cell(1, k + 1).paragraphs[0].add_run(v)
        table.cell(1, k + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, k + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    for k, v in enumerate(city):
        table_run = table.cell(k + 2, 0).paragraphs[0].add_run(v)
        table.cell(k + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    # 各地市班车（辆）、包车（辆）、危货（辆）、合计（辆）
    for k, v in enumerate(city):
        for i in range(4):
            if i != 3:
                number = '-' if TrajectoryIntegrityRate_dict[v][i] == 0 else TrajectoryIntegrityRate_dict[v][i]
                table_run = table.cell(k + 2, i + 1).paragraphs[0].add_run(str(number))
                table.cell(k + 2, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)
            else:
                number = '-' if TrajectoryIntegrityRate_dict[v][0] + TrajectoryIntegrityRate_dict[v][1] + \
                                TrajectoryIntegrityRate_dict[v][2] == 0 else TrajectoryIntegrityRate_dict[v][0] + \
                                                                             TrajectoryIntegrityRate_dict[v][1] + \
                                                                             TrajectoryIntegrityRate_dict[v][2]
                table_run = table.cell(k + 2, 4).paragraphs[0].add_run(str(number))
                table.cell(k + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)

    # 小计（辆）
    sum = 0
    for i in range(4):
        sum = 0
        if i != 3:
            num = [TrajectoryIntegrityRate_dict[j][i] for j in TrajectoryIntegrityRate_dict]
            for j in num:
                sum += j
            table_run = table.cell(17, i + 1).paragraphs[0].add_run(str(sum))
            table.cell(17, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(17, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
        else:
            for j in TrajectoryIntegrityRate_dict:
                for k in range(3):
                    sum += TrajectoryIntegrityRate_dict[j][k]
            table_run = table.cell(18, 1).paragraphs[0].add_run(str(sum))
            table.cell(18, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
            # 备注数量
            count = select_db(
                'SELECT count(*)  FROM Ve_Status_Details s  WHERE  s.行驶里程_公里  <> 0 and s.轨迹完整率 < 95 ')
            table_run = table.cell(19, 0).paragraphs[0].add_run(
                '备注：轨迹完整率低于99%的车辆总数{}台，低于95%车辆总数{}台。'.format(sum, count[0][0]))
            table.cell(19, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(19, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = '仿宋_GB2312'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            table_run.font.size = Pt(12)
    # else:
    #     table_run = table.cell(17, 4).paragraphs[0].add_run('-')
    #     table.cell(17, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     table.cell(17, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #     table_run.font.name = 'Times New Roman'
    #     table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    #     table_run.font.size = Pt(12)

    # 排名
    for k, v in enumerate(city):
        number = getSortByDict(TrajectoryIntegrityRate_dict, v)
        table_run = table.cell(k + 2, 5).paragraphs[0].add_run(str(number))
        table.cell(k + 2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    #################全省车辆漂移率问题清单########################################################################
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆漂移率问题清单 ')
    run.font.name = '方正小标宋简体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)

    # 添加word表格(Table Grid央视表格)
    table = doc.add_table(rows=20, cols=6, style='Table Grid')
    # 设置行高
    for i in range(20):
        table.rows[i].height = Cm(0.98)

    # 合并单元格
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(18, 1).merge(table.cell(18, 5))
    table.cell(19, 0).merge(table.cell(19, 5))

    # 在单元格中插入内容
    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 在单元格中插入内容
    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 4).paragraphs[0].add_run('合计（辆）')
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 5).paragraphs[0].add_run('排名')
    table.cell(0, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 0).paragraphs[0].add_run('小计(辆)')
    table.cell(17, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 0).paragraphs[0].add_run('总计(辆)')
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(17, 5).paragraphs[0].add_run('-')
    table.cell(17, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(17, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    content = ['班车(辆)', '包车(辆)', '危货(辆)']
    for k, v in enumerate(content):
        table_run = table.cell(1, k + 1).paragraphs[0].add_run(v)
        table.cell(1, k + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, k + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    for k, v in enumerate(city):
        table_run = table.cell(k + 2, 0).paragraphs[0].add_run(v)
        table.cell(k + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    # 各地市班车（辆）、包车（辆）、危货（辆）、合计（辆）
    for k, v in enumerate(city):
        for i in range(4):
            if i != 3:
                number = '-' if TrajectoryDriftRate_dict[v][i] == 0 else TrajectoryDriftRate_dict[v][i]
                table_run = table.cell(k + 2, i + 1).paragraphs[0].add_run(str(number))
                table.cell(k + 2, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)
            else:
                number = '-' if TrajectoryDriftRate_dict[v][0] + TrajectoryDriftRate_dict[v][1] + \
                                TrajectoryDriftRate_dict[v][2] == 0 else TrajectoryDriftRate_dict[v][0] + \
                                                                         TrajectoryDriftRate_dict[v][1] + \
                                                                         TrajectoryDriftRate_dict[v][2]
                table_run = table.cell(k + 2, 4).paragraphs[0].add_run(str(number))
                table.cell(k + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(k + 2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table_run.font.name = 'Times New Roman'
                table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                table_run.font.size = Pt(12)

    # 小计（辆）
    sum = 0
    for i in range(4):
        sum = 0
        if i != 3:
            num = [TrajectoryDriftRate_dict[j][i] for j in TrajectoryDriftRate_dict]
            for j in num:
                sum += j
            table_run = table.cell(17, i + 1).paragraphs[0].add_run(str(sum))
            table.cell(17, i + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(17, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
        else:
            for j in TrajectoryDriftRate_dict:
                for k in range(3):
                    sum += TrajectoryDriftRate_dict[j][k]
            table_run = table.cell(18, 1).paragraphs[0].add_run(str(sum))
            table.cell(18, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
            # 备注数量
            count = select_db(
                'SELECT count(*)  FROM Ve_Status_Details s  WHERE  s.轨迹漂移率 > 1  ')
            table_run = table.cell(19, 0).paragraphs[0].add_run(
                '备注：漂移率高于0.4%车辆总数{}台，高于1%以下的车辆总数{}台。'.format(sum, count[0][0]))
            table.cell(19, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(19, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = '仿宋_GB2312'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            table_run.font.size = Pt(12)
    else:
        table_run = table.cell(17, 4).paragraphs[0].add_run('-')
        table.cell(17, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(17, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 排名
    for k, v in enumerate(city):
        number = getSortByDict(TrajectoryDriftRate_dict, v)
        table_run = table.cell(k + 2, 5).paragraphs[0].add_run(str(number))
        table.cell(k + 2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)
    #################全省车辆超速报警问题清单########################################################################
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆超速报警问题清单')
    run.font.name = '方正小标宋简体'
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)

    # 添加word表格(Table Grid央视表格)
    table = doc.add_table(rows=21, cols=9, style='Table Grid')
    # 设置行高
    for i in range(21):
        table.rows[i].height = Cm(0.9)
    # 设置列宽
    table.cell(3, 0).width = Cm(2.2)
    table.cell(3, 1).width = Cm(2.2)
    table.cell(3, 3).width = Cm(2.2)
    table.cell(3, 5).width = Cm(2.2)
    table.cell(3, 7).width = Cm(2.2)
    table.cell(3, 2).width = Cm(1.5)
    table.cell(3, 4).width = Cm(1.5)
    table.cell(3, 6).width = Cm(1.5)
    table.cell(3, 8).width = Cm(1.5)

    # 合并单元格
    table.cell(0, 0).merge(table.cell(2, 0))
    table.cell(0, 1).merge(table.cell(0, 6))
    table.cell(1, 1).merge(table.cell(1, 2))
    table.cell(1, 3).merge(table.cell(1, 4))
    table.cell(1, 5).merge(table.cell(1, 6))
    table.cell(0, 7).merge(table.cell(1, 8))
    table.cell(20, 1).merge(table.cell(20, 8))
    table.cell(19, 1).merge(table.cell(19, 8))

    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 8).paragraphs[0].add_run('合计')
    table.cell(0, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(1, 1).paragraphs[0].add_run('班车')
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(1, 3).paragraphs[0].add_run('包车')
    table.cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(1, 5).paragraphs[0].add_run('危货')
    table.cell(1, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 1).paragraphs[0].add_run('报警次数(台次)')
    table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 3).paragraphs[0].add_run('报警次数（台次）')
    table.cell(2, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 5).paragraphs[0].add_run('报警次数（台次）')
    table.cell(2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 7).paragraphs[0].add_run('报警次数（台次）')
    table.cell(2, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 2).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 4).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 6).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 8).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 0).paragraphs[0].add_run('小计')
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(19, 0).paragraphs[0].add_run('总计(报警次数)')
    table.cell(19, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(19, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(20, 0).paragraphs[0].add_run('总计(车辆数)')
    table.cell(20, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(20, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 7).paragraphs[0].add_run('-')
    table.cell(18, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 8).paragraphs[0].add_run('-')
    table.cell(18, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    for k, v in enumerate(city):
        table_run = table.cell(k + 3, 0).paragraphs[0].add_run(v)
        table.cell(k + 3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 3, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    for k, v in enumerate(city):
        for i in range(3):
            number = '-' if Ve_Tired_Details_dict[v][i] == 0 else Ve_Tired_Details_dict[v][i]
            table_run = table.cell(k + 3, i * 2 + 2).paragraphs[0].add_run(str(number))
            table.cell(k + 3, i * 2 + 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(k + 3, i * 2 + 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)

    for k, v in enumerate(city):
        for i in range(3):
            number = '-' if Ve_Tired_Details_dict[v][i + 3] == 0 else Ve_Tired_Details_dict[v][i + 3]
            table_run = table.cell(k + 3, i * 2 + 1).paragraphs[0].add_run(str(number))
            table.cell(k + 3, i * 2 + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(k + 3, i * 2 + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)

    # 合计-报警次数（台次）
    for k, v in enumerate(city):
        number = '-' if Ve_Tired_Details_dict[v][0] + Ve_Tired_Details_dict[v][1] + Ve_Tired_Details_dict[v][
            2] == 0 else Ve_Tired_Details_dict[v][0] + Ve_Tired_Details_dict[v][1] + Ve_Tired_Details_dict[v][2]
        table_run = table.cell(k + 3, 8).paragraphs[0].add_run(str(number))
        table.cell(k + 3, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 3, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 合计-车辆数（辆）
    for k, v in enumerate(city):
        number = '-' if Ve_Tired_Details_dict[v][3] + Ve_Tired_Details_dict[v][4] + Ve_Tired_Details_dict[v][
            5] == 0 else Ve_Tired_Details_dict[v][3] + Ve_Tired_Details_dict[v][4] + Ve_Tired_Details_dict[v][5]
        table_run = table.cell(k + 3, 7).paragraphs[0].add_run(str(number))
        table.cell(k + 3, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 3, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 小计
    for i in range(3):
        sum = 0
        for k, v in enumerate(city):
            sum += Ve_Tired_Details_dict[v][i]
        else:
            table_run = table.cell(18, i * 2 + 2).paragraphs[0].add_run(str(sum))
            table.cell(18, i * 2 + 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, i * 2 + 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
    for i in range(3):
        sum = 0
        for k, v in enumerate(city):
            sum += Ve_Tired_Details_dict[v][i + 3]
        else:
            table_run = table.cell(18, i * 2 + 1).paragraphs[0].add_run(str(sum))
            table.cell(18, i * 2 + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, i * 2 + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)

    # 总计(报警次数)
    number = select_db('SELECT ifnull(sum(s.超速报警数),0) FROM Ve_Tired_Details s WHERE    s.超速报警数<> 0 ')
    table_run = table.cell(19, 1).paragraphs[0].add_run(str(number[0][0]))
    table.cell(19, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(19, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    table_run.font.size = Pt(12)

    # 总计(车辆数)
    number = select_db('SELECT count(*) FROM Ve_Tired_Details s WHERE    s.超速报警数<> 0  ')
    table_run = table.cell(20, 1).paragraphs[0].add_run(str(number[0][0]))
    table.cell(20, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(20, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    table_run.font.size = Pt(12)

    #################全省车辆疲劳驾驶报警问题清单########################################################################

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('全省车辆疲劳驾驶报警问题清单')
    run.font.name = '方正小标宋简体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run('统计周期:{}-{}'.format(tadaySevenAgoTime, tadayTime))
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = True
    run.font.size = Pt(12)
    paragraph1.space_after = Pt(4)

    # 添加word表格(Table Grid央视表格)
    table = doc.add_table(rows=21, cols=9, style='Table Grid')
    # 设置行高
    for i in range(21):
        table.rows[i].height = Cm(0.88)
    # 设置列宽
    table.cell(3, 0).width = Cm(2.2)
    table.cell(3, 1).width = Cm(2.2)
    table.cell(3, 3).width = Cm(2.2)
    table.cell(3, 5).width = Cm(2.2)
    table.cell(3, 7).width = Cm(2.2)
    # table.cell(3, 2).width = Cm(1.5)
    # table.cell(3, 4).width = Cm(1.5)
    # table.cell(3, 6).width = Cm(1.5)
    # table.cell(3, 8).width = Cm(1.5)

    # table.cell(2, 7).width = Cm(2.1)
    # 合并单元格
    table.cell(0, 0).merge(table.cell(2, 0))
    table.cell(0, 1).merge(table.cell(0, 6))
    table.cell(1, 1).merge(table.cell(1, 2))
    table.cell(1, 3).merge(table.cell(1, 4))
    table.cell(1, 5).merge(table.cell(1, 6))
    table.cell(0, 7).merge(table.cell(1, 8))
    table.cell(20, 1).merge(table.cell(20, 8))
    table.cell(19, 1).merge(table.cell(19, 8))

    table_run = table.cell(0, 0).paragraphs[0].add_run('地区')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 1).paragraphs[0].add_run('类别')
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(0, 8).paragraphs[0].add_run('合计')
    table.cell(0, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(1, 1).paragraphs[0].add_run('班车')
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(1, 3).paragraphs[0].add_run('包车')
    table.cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(1, 5).paragraphs[0].add_run('危货')
    table.cell(1, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 1).paragraphs[0].add_run('报警次数(台次)')
    table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 3).paragraphs[0].add_run('报警次数（台次）')
    table.cell(2, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 5).paragraphs[0].add_run('报警次数（台次）')
    table.cell(2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 7).paragraphs[0].add_run('报警次数（台次）')
    table.cell(2, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 2).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 4).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 6).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(2, 8).paragraphs[0].add_run('车辆数（辆）')
    table.cell(2, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 0).paragraphs[0].add_run('小计')
    table.cell(18, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(19, 0).paragraphs[0].add_run('总计(报警次数)')
    table.cell(19, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(19, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(20, 0).paragraphs[0].add_run('总计(车辆数)')
    table.cell(20, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(20, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = '仿宋_GB2312'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 7).paragraphs[0].add_run('-')
    table.cell(18, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    table_run = table.cell(18, 8).paragraphs[0].add_run('-')
    table.cell(18, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(18, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    table_run.font.size = Pt(12)

    # 添加地市名称
    for k, v in enumerate(city):
        table_run = table.cell(k + 3, 0).paragraphs[0].add_run(v)
        table.cell(k + 3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 3, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = '仿宋_GB2312'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        table_run.font.size = Pt(12)

    # 各地市报警次数(台次)
    for k, v in enumerate(city):
        for i in range(3):
            number = '-' if Ve_Tired_Details_dict2[v][i] == 0 else Ve_Tired_Details_dict2[v][i]
            table_run = table.cell(k + 3, i * 2 + 2).paragraphs[0].add_run(str(number))
            table.cell(k + 3, i * 2 + 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(k + 3, i * 2 + 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)

    # 各地市车辆数（辆）
    for k, v in enumerate(city):
        for i in range(3):
            number = '-' if Ve_Tired_Details_dict2[v][i + 3] == 0 else Ve_Tired_Details_dict2[v][i + 3]
            table_run = table.cell(k + 3, i * 2 + 1).paragraphs[0].add_run(str(number))
            table.cell(k + 3, i * 2 + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(k + 3, i * 2 + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)

    # 合计-报警次数（台次）
    for k, v in enumerate(city):
        number = '-' if Ve_Tired_Details_dict2[v][0] + Ve_Tired_Details_dict2[v][1] + Ve_Tired_Details_dict2[v][
            2] == 0 else Ve_Tired_Details_dict2[v][0] + Ve_Tired_Details_dict2[v][1] + Ve_Tired_Details_dict2[v][2]
        table_run = table.cell(k + 3, 8).paragraphs[0].add_run(str(number))
        table.cell(k + 3, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 3, 8).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 合计-车辆数（辆）
    for k, v in enumerate(city):
        number = '-' if Ve_Tired_Details_dict2[v][3] + Ve_Tired_Details_dict2[v][4] + Ve_Tired_Details_dict2[v][
            5] == 0 else Ve_Tired_Details_dict2[v][3] + Ve_Tired_Details_dict2[v][4] + Ve_Tired_Details_dict2[v][5]
        table_run = table.cell(k + 3, 7).paragraphs[0].add_run(str(number))
        table.cell(k + 3, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(k + 3, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table_run.font.name = 'Times New Roman'
        table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        table_run.font.size = Pt(12)

    # 小计
    for i in range(3):
        sum = 0
        for k, v in enumerate(city):
            # if v == '营口市' or v == '行政审批局':
            #     continue
            sum += Ve_Tired_Details_dict2[v][i]
        else:
            table_run = table.cell(18, i * 2 + 2).paragraphs[0].add_run(str(sum))
            table.cell(18, i * 2 + 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, i * 2 + 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)
    for i in range(3):
        sum = 0
        for k, v in enumerate(city):
            # if v == '营口市' or v == '行政审批局':
            #     continue
            sum += Ve_Tired_Details_dict2[v][i + 3]
        else:
            table_run = table.cell(18, i * 2 + 1).paragraphs[0].add_run(str(sum))
            table.cell(18, i * 2 + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(18, i * 2 + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_run.font.name = 'Times New Roman'
            table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            table_run.font.size = Pt(12)

    # 总计(报警次数)
    number = select_db('SELECT ifnull(sum(s.疲劳驾驶报警数),0) FROM Ve_Tired_Details s WHERE    s.疲劳驾驶报警数<> 0 ')
    table_run = table.cell(19, 1).paragraphs[0].add_run(str(number[0][0]))
    table.cell(19, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(19, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    table_run.font.size = Pt(12)

    # 总计(车辆数)
    number = select_db('SELECT count(*) FROM Ve_Tired_Details s WHERE    s.疲劳驾驶报警数<> 0  ')
    table_run = table.cell(20, 1).paragraphs[0].add_run(str(number[0][0]))
    table.cell(20, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(20, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = 'Times New Roman'
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    table_run.font.size = Pt(12)

    doc.save(path)


# 路径名称(D:\专项整治\专项整治2020年12月04日-2020年12月10日\1-问题清单2020年12月04日-2020年12月10日)
fileName = "专项整治{}-{}".format(
    oneWeekDayAgo.strftime('%Y') + '年' + oneWeekDayAgo.strftime('%m') + '月' + oneWeekDayAgo.strftime('%d') + '日',
    now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日')
fileName2 = "1-问题清单{}-{}".format(
    oneWeekDayAgo.strftime('%Y') + '年' + oneWeekDayAgo.strftime('%m') + '月' + oneWeekDayAgo.strftime('%d') + '日',
    now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日')

# 创建文件生成文件夹
path = "D://专项整治//{}//{}".format(fileName, fileName2, )
if not os.path.exists(path):
    os.makedirs(path)

createWord(now, oneWeekDayAgo, fileName, fileName2)
print('生成"问题清单.docx"完毕')

end_t = datetime.datetime.now()
finish_time2 = str((end_t - start_t).seconds)

print("------------------")
print("各地市EXCEL表执行了：" + finish_time + "秒")
print("问题清单执行了：" + finish_time2 + "秒")
print("------------------")

def getDesktopOfWindows():
    '''
    :return: 返回当前桌面路径(例 d://Users//Administrator//Desktop)
    '''
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    return winreg.QueryValueEx(key, "Desktop")[0].replace('\\','//')

def makeZipOfFolder(source_dir, output_filename):
    '''
    打包、压缩文件夹
    :param source_dir: 原文件夹
    :param output_filename: 压缩目标文件
    :return:
    '''
    zipf = zipfile.ZipFile(output_filename, 'w', zipfile.ZIP_DEFLATED)
    '''ZIP_DEFLATED:表示压缩，ZIP_STORE：表示只打包，不压缩'''
    pre_len = len(os.path.dirname(source_dir))
    for parent, dirnames, filenames in os.walk(source_dir):
        for filename in filenames:
          pathfile = os.path.join(parent, filename)
          arcname = pathfile[pre_len:].strip(os.path.sep)
          zipf.write(pathfile, '//'+arcname)
    zipf.close()

# 打包并压缩明细文件

zipName = "全省重点营运车辆联网联控专项整治行动车辆问题清单及车辆问题明细({}-{})".format(oneWeekDayAgo.strftime('%Y') + '年' + oneWeekDayAgo.strftime('%m') + '月' + oneWeekDayAgo.strftime('%d') + '日',now.strftime('%Y') + '年' + now.strftime('%m') + '月' + now.strftime('%d') + '日')
output_filename = '{}//{}.zip'.format(getDesktopOfWindows(),zipName)
source_dir = 'D://专项整治//{}'.format(fileName)
makeZipOfFolder(source_dir,output_filename)


print("脚本执行完毕。")

sg.popup("脚本执行完毕。",keep_on_top=True)