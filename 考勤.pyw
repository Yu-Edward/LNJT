from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落对齐样式
from docx.enum.table import WD_ALIGN_VERTICAL  # 表格垂直对齐方式
#from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt  # 段落字体磅数
#from docx.shared import RGBColor  # 字体颜色
from docx.shared import Inches  # 图标尺寸
from docx.shared import Cm  # 表格高度、宽度等(厘米 )
from docx.oxml.ns import qn  # 设置中文格式
import datetime,time,os,winreg,random,chinese_calendar
import PySimpleGUI as sg
from dateutil.relativedelta import relativedelta

'''上月26号到本月25号的考勤统计'''

# 上个月的月份和年份
lastMonth = datetime.datetime.strftime(datetime.date.today() - relativedelta(months=+1),'%m')
lastYear = datetime.datetime.strftime(datetime.date.today() - relativedelta(months=+1),'%Y')
# 当月的月份和年份
month = datetime.datetime.strftime(datetime.date.today(),'%m')
year = datetime.datetime.strftime(datetime.date.today(),'%Y')


datestart = datetime.datetime(int(lastYear),int(lastMonth),26)
days = []
dateend = datetime.datetime(int(year),int(month),25)
# while datestart <= dateend:
#     if datestart.isoweekday() != 6 and datestart.isoweekday() != 7:
#         days.append((datestart.month,datestart.day))
#     datestart += datetime.timedelta(days=1)

while datestart <= dateend:
    days.append((datestart.month,datestart.day)) if chinese_calendar.is_workday(datestart) else None
    datestart += datetime.timedelta(days=1)


# print('{}年{}月{}日-{}年{}月{}日,工作日共 {} 天。'.format(lastYear,lastMonth,26,year,month,25,len(days)))

def getDesktopOfWindows():
    '''
    :return: 返回当前桌面路径(例 d://Users//Administrator//Desktop)
    '''
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    return winreg.QueryValueEx(key, "Desktop")[0].replace('\\','//')


path = '{}//{}年{}{}'.format(getDesktopOfWindows(),year,month,'月驻省人员考勤表.docx')

doc = Document()


# 西文字体
doc.styles['Normal'].font.name = u'宋体'
# 中文字体
doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)

# 记录当前年、月
year = datetime.datetime.now().year
month = datetime.datetime.now().month

paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = paragraph.add_run(str(year))
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run.font.bold = True
run.font.size = Pt(22)
run = paragraph.add_run('年')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.bold = True
run.font.size = Pt(22)
run = paragraph.add_run(str(month))
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run.font.bold = True
run.font.size = Pt(22)
run = paragraph.add_run('月驻省人员（朱赛男）考勤表\r')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.bold = True
run.font.size = Pt(22)



table = doc.add_table(rows=len(days)+2, cols=6, style='Table Grid')

for i in range(len(days)+2):
    table.rows[i].height = Cm(0.5)

table.cell(0, 0).merge(table.cell(0, 1))
table.cell(0, 2).merge(table.cell(1, 2))
table.cell(0, 3).merge(table.cell(1, 3))
table.cell(0, 4).merge(table.cell(1, 4))
table.cell(0, 5).merge(table.cell(1, 5))

def insertContentIntoCell(rows,cols,size,font='宋体',text='None',bold=False):
    table_run = table.cell(rows, cols).paragraphs[0].add_run(text)
    table.cell(rows, cols).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(rows, cols).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = font
    table_run.font.bold = bold
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    table_run.font.size = Pt(size)

table.cell(2, 0).width = Inches(0.5)
table.cell(2, 1).width = Inches(0.5)
table.cell(2, 2).width = Inches(1)
table.cell(2, 3).width = Inches(1)
table.cell(2, 4).width = Inches(1.4)
table.cell(2, 5).width = Inches(1.6)


insertContentIntoCell(0,0,12,text='时间',bold=True)
insertContentIntoCell(1,0,12,text='月',bold=True)
insertContentIntoCell(1,1,12,text='日',bold=True)
insertContentIntoCell(0,2,12,text='上班时间',bold=True)
insertContentIntoCell(1,3,12,text='下班时间',bold=True)
insertContentIntoCell(1,4,12,text='是否邮件请假',bold=True)
insertContentIntoCell(1,5,12,text='未按时打卡说明',bold=True)

# for i in range(len(days)):
#     goWorkTime1 = random.randint(7, 8)
#     goWorkTime = '{}:{:>02d}'.format(goWorkTime1, random.randint(50, 59) if goWorkTime1 == 7 else random.randint(0, 25))
#     insertContentIntoCell(i+2,2,12,text=goWorkTime,bold=False)
# for i in range(len(days)):
#     goWorkTime = '{}:{:>02d}'.format(5, random.randint(35, 50))
#     insertContentIntoCell(i+2,3,12,text=goWorkTime,bold=False)

[insertContentIntoCell(i+2, 0, 12,text=str(days[i][0])) for i in range(len(days))]
[insertContentIntoCell(i+2, 1, 12,text=str(days[i][1])) for i in range(len(days))]


paragraph = doc.add_paragraph()
run = paragraph.add_run('\r\r\r')

paragraph = doc.add_paragraph()
paragraph.paragraph_format.first_line_indent = Cm(5.74)
run = paragraph.add_run('业主方确认签名：\r')
run.font.name = '宋体 (中文正文)'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体 (中文正文)')
run.font.bold = False
run.font.size = Pt(15)
paragraph = doc.add_paragraph()
paragraph.paragraph_format.first_line_indent = Cm(6.74)
run = paragraph.add_run('确认时间：')
run.font.name = '宋体 (中文正文)'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体 (中文正文)')
run.font.bold = False
run.font.size = Pt(15)

doc.add_page_break()
##################################################################################################

paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = paragraph.add_run(str(year))
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run.font.bold = True
run.font.size = Pt(22)
run = paragraph.add_run('年')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.bold = True
run.font.size = Pt(22)
run = paragraph.add_run(str(month))
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run.font.bold = True
run.font.size = Pt(22)
run = paragraph.add_run('月驻省人员（于罡）考勤表\r')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.bold = True
run.font.size = Pt(22)



table = doc.add_table(rows=len(days)+2, cols=6, style='Table Grid')

for i in range(len(days)+2):
    table.rows[i].height = Cm(0.5)

table.cell(0, 0).merge(table.cell(0, 1))
table.cell(0, 2).merge(table.cell(1, 2))
table.cell(0, 3).merge(table.cell(1, 3))
table.cell(0, 4).merge(table.cell(1, 4))
table.cell(0, 5).merge(table.cell(1, 5))

def insertContentIntoCell(rows,cols,size,font='宋体',text='None',bold=False):
    table_run = table.cell(rows, cols).paragraphs[0].add_run(text)
    table.cell(rows, cols).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(rows, cols).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_run.font.name = font
    table_run.font.bold = bold
    table_run.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    table_run.font.size = Pt(size)

table.cell(2, 0).width = Inches(0.5)
table.cell(2, 1).width = Inches(0.5)
table.cell(2, 2).width = Inches(1)
table.cell(2, 3).width = Inches(1)
table.cell(2, 4).width = Inches(1.4)
table.cell(2, 5).width = Inches(1.6)


insertContentIntoCell(0,0,12,text='时间',bold=True)
insertContentIntoCell(1,0,12,text='月',bold=True)
insertContentIntoCell(1,1,12,text='日',bold=True)
insertContentIntoCell(0,2,12,text='上班时间',bold=True)
insertContentIntoCell(1,3,12,text='下班时间',bold=True)
insertContentIntoCell(1,4,12,text='是否邮件请假',bold=True)
insertContentIntoCell(1,5,12,text='未按时打卡说明',bold=True)

[insertContentIntoCell(i+2, 0, 12,text=str(days[i][0])) for i in range(len(days))]
[insertContentIntoCell(i+2, 1, 12,text=str(days[i][1])) for i in range(len(days))]


paragraph = doc.add_paragraph()
run = paragraph.add_run('\r\r\r')

paragraph = doc.add_paragraph()
paragraph.paragraph_format.first_line_indent = Cm(5.74)
run = paragraph.add_run('业主方确认签名：\r')
run.font.name = '宋体 (中文正文)'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体 (中文正文)')
run.font.bold = False
run.font.size = Pt(15)
paragraph = doc.add_paragraph()
paragraph.paragraph_format.first_line_indent = Cm(6.74)
run = paragraph.add_run('确认时间：')
run.font.name = '宋体 (中文正文)'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体 (中文正文)')
run.font.bold = False
run.font.size = Pt(15)





doc.save(path)


sg.popup('{}年{}月{}日-{}年{}月{}日,工作日共 {} 天。'.format(lastYear,lastMonth,26,year,month,25,len(days)),keep_on_top = True)

# sg.popup('Finish')